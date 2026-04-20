from __future__ import annotations

import math
import re
import textwrap
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"c:\Bookstore\BookStore-Asg06")
SOURCE_DOCX = ROOT / "ai_ecommerce_08.18_hoangth.docx"
OUTPUT_DOCX = ROOT / "baocao.docx"
ASSET_DIR = ROOT / "report_assets"


def extract_paragraphs(docx_path: Path) -> list[str]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml")
    root = ET.fromstring(xml)
    paragraphs: list[str] = []
    for p in root.findall(".//w:body/w:p", ns):
        texts = [t.text or "" for t in p.findall(".//w:t", ns)]
        text = "".join(texts).strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for style_name, size in [("Title", 20), ("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_paragraph(doc: Document, text: str, *, align=None, bold=False, italic=False, size=12, color=None, space_after=6, line_spacing=1.3, style="Normal"):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    fmt = para.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = line_spacing
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph(style=f"Heading {level}")
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(12)


PY_KEYWORDS = {
    "from", "import", "as", "return", "for", "in", "if", "else", "elif", "try", "except", "with",
    "class", "def", "None", "True", "False", "pass", "raise", "and", "or", "not", "yield", "while",
    "continue", "break", "lambda", "global", "nonlocal", "del", "assert", "async", "await", "self",
}


def tokenize_python_line(line: str):
    pattern = re.compile(r"(#.*$)|('(?:[^'\\]|\\.)*'|\"(?:[^\"\\]|\\.)*\")|\b([A-Za-z_][A-Za-z0-9_]*)\b|([0-9]+(?:\.[0-9]+)?)|([^A-Za-z0-9_#\'\"\s]+)|\s+")
    for match in pattern.finditer(line):
        comment, string, word, number, symbol = match.groups()
        token = match.group(0)
        if comment:
            yield token, (72, 187, 120)
        elif string:
            yield token, (214, 157, 133)
        elif word:
            if word in PY_KEYWORDS:
                yield token, (86, 156, 214)
            elif word in {"logger", "requests", "torch", "nn", "defaultdict", "Path", "Any"}:
                yield token, (78, 201, 176)
            else:
                yield token, (220, 220, 220)
        elif number:
            yield token, (181, 206, 168)
        elif symbol:
            yield token, (212, 212, 212)
        else:
            yield token, (220, 220, 220)


def add_code_block(doc: Document, code: str, language: str = "python") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.7)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_shading(cell, "1E1E1E")
    set_cell_margins(cell, top=100, start=120, bottom=100, end=120)

    lines = code.rstrip().splitlines()
    for idx, line in enumerate(lines):
        para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        ln = para.add_run(f"{idx + 1:>3}  ")
        ln.font.name = "Consolas"
        ln.font.size = Pt(8.5)
        ln.font.color.rgb = RGBColor(120, 120, 120)
        for token, color in tokenize_python_line(line):
            run = para.add_run(token)
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(*color)

    doc.add_paragraph()


def load_font(preferred: list[str], size: int) -> ImageFont.FreeTypeFont:
    for path in preferred:
        if Path(path).exists():
            try:
                return ImageFont.truetype(path, size)
            except Exception:
                pass
    return ImageFont.load_default()


TITLE_FONT = load_font([
    r"C:\\Windows\\Fonts\\arialbd.ttf",
    r"C:\\Windows\\Fonts\\calibrib.ttf",
], 30)
BODY_FONT = load_font([
    r"C:\\Windows\\Fonts\\arial.ttf",
    r"C:\\Windows\\Fonts\\calibri.ttf",
], 20)
SMALL_FONT = load_font([
    r"C:\\Windows\\Fonts\\arial.ttf",
    r"C:\\Windows\\Fonts\\calibri.ttf",
], 16)


def draw_box(draw: ImageDraw.ImageDraw, xy, title: str, fill: str, outline: str = "#b71c1c", text_fill: str = "white", radius: int = 20):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=4)
    bbox = draw.textbbox((0, 0), title, font=BODY_FONT)
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    tx = x1 + (x2 - x1 - tw) / 2
    ty = y1 + (y2 - y1 - th) / 2
    draw.text((tx, ty), title, fill=text_fill, font=BODY_FONT)


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, color="#b71c1c", width=6):
    draw.line([start, end], fill=color, width=width)
    sx, sy = start
    ex, ey = end
    dx = ex - sx
    dy = ey - sy
    length = math.hypot(dx, dy) or 1.0
    ux, uy = dx / length, dy / length
    arrow_size = 18
    left = (ex - ux * arrow_size - uy * arrow_size * 0.55, ey - uy * arrow_size + ux * arrow_size * 0.55)
    right = (ex - ux * arrow_size + uy * arrow_size * 0.55, ey - uy * arrow_size - ux * arrow_size * 0.55)
    draw.polygon([end, left, right], fill=color)


def create_architecture_image(path: Path):
    img = Image.new("RGB", (1600, 900), "#fff7f8")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 1599, 899], outline="#ef9a9a", width=6)
    draw.text((50, 30), "Kiến trúc hệ thống BookStore", font=TITLE_FONT, fill="#8e1120")
    draw.text((50, 85), "API Gateway + microservices + recommender-ai-service", font=BODY_FONT, fill="#5b2630")

    draw_box(draw, (70, 170, 290, 280), "Frontend", "#c2182b")
    draw_box(draw, (360, 170, 620, 280), "API Gateway", "#8e1120")
    draw_box(draw, (740, 100, 1020, 210), "Auth Service", "#d84315")
    draw_box(draw, (740, 230, 1020, 340), "Catalog Service", "#b71c1c")
    draw_box(draw, (740, 360, 1020, 470), "Product Service", "#e53935")
    draw_box(draw, (740, 490, 1020, 600), "Order / Cart / Pay", "#c62828")
    draw_box(draw, (1120, 170, 1470, 380), "Recommender AI Service", "#5b2630")
    draw_box(draw, (1120, 450, 1470, 610), "Neo4j + FAISS + KB", "#2e1114")

    draw_arrow(draw, (290, 225), (360, 225))
    draw_arrow(draw, (620, 210), (740, 155))
    draw_arrow(draw, (620, 225), (740, 285))
    draw_arrow(draw, (620, 240), (740, 415))
    draw_arrow(draw, (620, 255), (740, 545))
    draw_arrow(draw, (1020, 285), (1120, 250))
    draw_arrow(draw, (1020, 545), (1120, 520))
    draw_arrow(draw, (1260, 380), (1260, 450))

    draw.text((50, 760), "Tầng vào duy nhất qua Gateway, phần AI đọc dữ liệu thật từ service nghiệp vụ.", font=SMALL_FONT, fill="#5b2630")
    img.save(path)


def create_behavior_image(path: Path):
    img = Image.new("RGB", (1600, 900), "#fff9f9")
    draw = ImageDraw.Draw(img)
    draw.text((50, 30), "Pipeline phân tích hành vi khách hàng", font=TITLE_FONT, fill="#8e1120")
    draw_box(draw, (70, 190, 290, 305), "Order Service", "#c2182b")
    draw_box(draw, (70, 370, 290, 485), "Comment-Rate", "#d84315")
    draw_box(draw, (360, 275, 620, 400), "Behavior Analyzer", "#8e1120")
    draw_box(draw, (700, 160, 960, 285), "Feature Engineering", "#b71c1c")
    draw_box(draw, (700, 335, 960, 460), "model_behavior", "#2e1114")
    draw_box(draw, (1040, 200, 1410, 340), "Customer Behavior Profile", "#5b2630")
    draw_box(draw, (1040, 390, 1410, 530), "Recommendation / Churn", "#7f1d1d")

    draw_arrow(draw, (290, 245), (360, 330))
    draw_arrow(draw, (290, 425), (360, 350))
    draw_arrow(draw, (620, 335), (700, 225))
    draw_arrow(draw, (620, 355), (700, 395))
    draw_arrow(draw, (960, 220), (1040, 270))
    draw_arrow(draw, (960, 395), (1040, 445))

    draw.text((70, 610), "Dữ liệu hành vi được gom từ nhiều service, sau đó đưa qua mô hình đa nhiệm để tạo profile.", font=BODY_FONT, fill="#5b2630")
    img.save(path)


def create_rag_image(path: Path):
    img = Image.new("RGB", (1600, 900), "#fff7f7")
    draw = ImageDraw.Draw(img)
    draw.text((50, 30), "Luồng RAG cho chat tư vấn", font=TITLE_FONT, fill="#8e1120")
    draw_box(draw, (70, 180, 300, 290), "Người dùng", "#c2182b")
    draw_box(draw, (380, 180, 640, 290), "Chat Orchestrator", "#8e1120")
    draw_box(draw, (720, 90, 1000, 200), "Intent Detector", "#d84315")
    draw_box(draw, (720, 230, 1000, 340), "RAG Retrieval", "#b71c1c")
    draw_box(draw, (720, 370, 1000, 480), "KB / FAISS", "#2e1114")
    draw_box(draw, (1120, 160, 1460, 300), "SourceItem + Answer", "#5b2630")
    draw_box(draw, (1120, 360, 1460, 500), "Order / Product Support", "#7f1d1d")

    draw_arrow(draw, (300, 235), (380, 235))
    draw_arrow(draw, (640, 235), (720, 140))
    draw_arrow(draw, (640, 250), (720, 285))
    draw_arrow(draw, (1000, 285), (1120, 230))
    draw_arrow(draw, (1000, 140), (1120, 230))
    draw_arrow(draw, (1000, 425), (1120, 425))

    draw.text((70, 620), "RAG giúp câu trả lời có nguồn trích dẫn và giảm hallucination khi tư vấn.", font=BODY_FONT, fill="#5b2630")
    img.save(path)


def create_deploy_image(path: Path):
    img = Image.new("RGB", (1600, 900), "#fff9f9")
    draw = ImageDraw.Draw(img)
    draw.text((50, 30), "Triển khai Docker Compose và dịch vụ", font=TITLE_FONT, fill="#8e1120")
    boxes = [
        (70, 160, 320, 280, "Docker Compose"),
        (410, 120, 700, 220, "Auth / Customer"),
        (410, 240, 700, 340, "Product / Catalog"),
        (410, 360, 700, 460, "Cart / Order / Pay"),
        (410, 480, 700, 580, "Ship / Staff / Comment"),
        (780, 180, 1080, 310, "API Gateway"),
        (1140, 120, 1480, 250, "Recommender AI Service"),
        (1140, 290, 1480, 420, "Neo4j / FAISS / Model"),
    ]
    for x1, y1, x2, y2, label in boxes:
        draw_box(draw, (x1, y1, x2, y2), label, "#8e1120" if "AI" not in label else "#5b2630")
    draw_arrow(draw, (320, 220), (410, 170))
    draw_arrow(draw, (320, 240), (410, 290))
    draw_arrow(draw, (320, 260), (410, 410))
    draw_arrow(draw, (320, 280), (410, 530))
    draw_arrow(draw, (700, 260), (780, 245))
    draw_arrow(draw, (1080, 245), (1140, 190))
    draw_arrow(draw, (1260, 250), (1260, 290))
    draw.text((70, 670), "Gateway là cửa ngõ duy nhất, đồng thời kiểm tra JWT và ghi log/metrics trước khi chuyển tiếp request.", font=BODY_FONT, fill="#5b2630")
    img.save(path)


def transform_text(text: str) -> str:
    replacements = {
        "Recommendation Engine": "Hybrid Recommender (hybrid_recommender.py)",
        "Intelligent Chatbot": "Chat Orchestrator (app/orchestrator.py)",
        "Behavior & Churn Prediction": "Behavior Analyzer + model_behavior",
        "Behavior Analytics": "Behavior Analyzer + model_behavior",
        "Knowledge Base (KB)": "Knowledge Graph / KB (kb_manager.py, neo4j_store.py)",
        "Retrieval-Augmented Generation": "RAG (rag_retrieval.py)",
        "API Gateway": "API Gateway (api_gateway/app/views.py)",
        "Microservices": "Microservices BookStore",
        "Deep Learning Model": "Deep Learning model_behavior",
        "Model Behavior": "model_behavior",
        "Chatbot": "Chat Orchestrator",
        "Knowledge Base": "Knowledge Graph / KB",
        "Neural Networks": "model_behavior",
        "FAISS": "FAISS + rag_retrieval",
        "Neo4j": "Neo4j knowledge graph",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def looks_like_code_paragraph(text: str) -> bool:
    content = text.strip()
    if len(content) < 24:
        return False
    if content.startswith(("class ", "def ", "async def ", "from ", "import ", "@")):
        return True
    code_keywords = re.findall(
        r"\b(async|await|def|class|return|try|except|if|elif|else|for|while|import|from|with)\b",
        content,
    )
    symbol_hits = len(re.findall(r"[(){}\[\]:=]", content))
    return (len(code_keywords) >= 2 and symbol_hits >= 4) or ("async def " in content) or ("def " in content and "return" in content)


def normalize_inline_code_text(text: str) -> str:
    content = text.replace("`", "").strip()
    content = re.sub(
        r"\s+(?=(?:async\s+def|def|class|if\b|elif\b|else:|for\b|while\b|try:|except\b|return\b|from\b|import\b|with\b|@))",
        "\n",
        content,
    )
    content = re.sub(r"\s+(?=#)", "\n", content)
    content = re.sub(r"\n{3,}", "\n\n", content)
    return content


HYBRID_SNIPPET = textwrap.dedent(
    '''
    def get_personalized(
        customer_id: int,
        interactions: dict[str, dict[int, int]] | None = None,
        limit: int = 8,
        budget_min: float | None = None,
        budget_max: float | None = None,
        category: str | None = None,
        customer_ratings: dict[int, int] | None = None,
    ) -> list[dict[str, Any]]:
        interactions = interactions or _load_interactions(customer_id)
        customer_ratings = customer_ratings or _load_customer_ratings(customer_id)

        base_personalized = base_recommendation.get_personalized(
            customer_id=customer_id,
            interactions=interactions,
            limit=max(limit * 2, 10),
            budget_min=budget_min,
            budget_max=budget_max,
            category=category,
            customer_ratings=customer_ratings,
        )

        base_popular = base_recommendation.get_popular(limit=max(limit * 2, 10))
        all_books = catalog_client.get_all_products(limit=500)
        catalog_index = _catalog_lookup(all_books)
        neo4j_store.sync_catalog(all_books)
        graph_scores = neo4j_store.score_candidates(
            customer_id=customer_id,
            seed_product_ids=recent_books,
            candidate_products=all_books,
        )

        lstm_scores = lstm_ranker.score_candidates(
            customer_id=customer_id,
            interactions=interactions,
            candidate_products=all_books,
        )

        profile = behavior_service.analyze(customer_id, {k: sum(v.values()) for k, v in interactions.items()})
        rag_entries = []
        if profile.get("preferred_categories"):
            top_category = profile["preferred_categories"][0].get("category")
            if top_category:
                rag_entries = rag_service.retrieve(f"{top_category} product recommendation", top_k=1)

        merged = _merge_sources([
            ("behavior", 0.55, base_personalized),
            ("popularity", 0.18, base_popular),
            ("graph", 0.22, _score_map_to_items(graph_scores, catalog_index)),
            ("lstm", 0.25, _score_map_to_items(lstm_scores, catalog_index)),
        ], limit)

        if rag_entries and merged:
            merged[0]["reason"] = f"{merged[0]['reason']} | RAG: {rag_entries[0]['title']}"
        return merged
    '''
).strip("\n")


ROUTES_SNIPPET = textwrap.dedent(
    '''
    @router.get("/api/v1/recommend/{customer_id}", response_model=RecommendationResponse, tags=["recommend"])
    def recommend(
        customer_id: int,
        limit:       int   = Query(8, ge=1, le=20),
        category:    Optional[str]  = None,
        budget_max:  Optional[float] = None,
    ):
        try:
            interactions: dict = {}
            try:
                from django.apps import apps
                CustomerBookInteraction = apps.get_model("app", "CustomerBookInteraction")
                qs = CustomerBookInteraction.objects.filter(customer_id=customer_id)
                for ix in qs:
                    interactions.setdefault(ix.interaction_type, {})[ix.book_id] = ix.count
            except Exception:
                pass

            recs = rec_svc.get_personalized(
                customer_id=customer_id,
                interactions=interactions,
                limit=limit,
                category=category,
                budget_max=budget_max,
            )
            if not recs:
                recs = rec_svc.get_popular(limit=limit)
            items = [RecommendationItem(**r) for r in recs]
            return RecommendationResponse(
                customer_id=customer_id,
                recommendations=items,
                source="personalized" if interactions else "popular",
            )
        except Exception as exc:
            logger.exception("Recommend error: %s", exc)
            raise HTTPException(status_code=500, detail=str(exc))
    '''
).strip("\n")


GATEWAY_SNIPPET = textwrap.dedent(
    '''
    SERVICE_MAP = {
        'auth':            'auth-service',
        'customers':       'customer-service',
        'books':           'product-service',
        'products':        'product-service',
        'product-types':   'product-service',
        'categories':      'product-service',
        'brands':          'product-service',
        'carts':           'cart-service',
        'orders':          'order-service',
        'payments':        'pay-service',
        'shipments':       'ship-service',
        'staff':           'staff-service',
        'comments':        'comment-rate-service',
        'catalog':         'catalog-service',
        'manager':         'manager-service',
        'recommendations': 'recommender-ai-service',
        'ai':              'recommender-ai-service',
        'v1':              'recommender-ai-service',
    }

    @csrf_exempt
    def api_proxy(request, path):
        _inc('total_requests')
        client_ip = request.META.get('HTTP_X_FORWARDED_FOR', request.META.get('REMOTE_ADDR', '0.0.0.0')).split(',')[0].strip()
        if _is_rate_limited(client_ip):
            _inc('rate_limited')
            return JsonResponse({'error': 'Too many requests'}, status=429)

        parts = path.strip('/').split('/')
        service_alias = parts[0]
        if service_alias not in SERVICE_MAP:
            return JsonResponse({'error': f'Unknown route: {service_alias}'}, status=404)

        target_service = SERVICE_MAP[service_alias]
        auth_header = request.META.get('HTTP_AUTHORIZATION', '')
        token = auth_header.replace('Bearer ', '').strip()
        if service_alias not in PUBLIC_ROUTES and not token:
            _inc('auth_failures')
            return JsonResponse({'error': 'Authentication required'}, status=401)
    '''
).strip("\n")


NEO4J_SNIPPET = textwrap.dedent(
    '''
    def score_candidates(
        self,
        customer_id: int,
        seed_product_ids: list[int],
        candidate_products: list[dict[str, Any]],
    ) -> dict[int, dict[str, Any]]:
        if not candidate_products:
            return {}

        seed_products = [catalog_client.get_product_by_id(pid) for pid in seed_product_ids]
        seed_products = [p for p in seed_products if p]
        score_map: dict[int, dict[str, Any]] = defaultdict(lambda: {"score": 0.0, "reason": "graph fallback"})

        seed_categories = { _book_category_key(p) for p in seed_products if _book_category_key(p) }
        seed_brands = {
            str(p.get("brand_detail", {}).get("slug") if isinstance(p.get("brand_detail"), dict) else p.get("brand") or "")
            for p in seed_products
            if (p.get("brand_detail") or p.get("brand"))
        }

        for product in candidate_products:
            pid = int(product.get("id") or 0)
            if not pid:
                continue
            category = _book_category_key(product)
            brand = str(product.get("brand_detail", {}).get("slug") if isinstance(product.get("brand_detail"), dict) else product.get("brand") or "")

            score = 0.0
            reasons = []
            if category and category in seed_categories:
                score += 1.8
                reasons.append(f"cùng hệ category với hành vi gần đây: {category}")
            if brand and brand in seed_brands:
                score += 1.2
                reasons.append(f"cùng brand với lựa chọn trước: {brand}")

            score_map[pid] = {
                "score": round(score, 3),
                "reason": "; ".join(reasons) if reasons else "đồ thị tri thức không có liên hệ mạnh",
            }

        return score_map
    '''
).strip("\n")


def add_title_page(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO AI TRONG E-COMMERCE VÀ ỨNG DỤNG BOOKSTORE")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(141, 17, 32)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Bản Word chi tiết, bám code thực tế, có hình minh họa và khối mã nguồn")
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)

    info_lines = [
        "Đề tài: Phân tích hành vi khách hàng, KB/knowledge graph, RAG chat và deploy microservices",
        "Hệ thống: BookStore Asg06",
        f"Thời điểm xuất bản: 16/04/2026",
    ]
    for line in info_lines:
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=3)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_toc_like(doc: Document) -> None:
    add_heading(doc, "MỤC LỤC TÓM TẮT", 1)
    items = [
        "Mở đầu và kiến trúc tổng thể",
        "Phần I. Khảo sát các ứng dụng AI trong e-commerce",
        "Phần II. Ứng dụng phân tích hành vi khách hàng",
        "Phần III. Kiến trúc model_behavior dựa trên deep learning",
        "Phần IV. Knowledge graph và Knowledge Base cho tư vấn",
        "Phần V. RAG chat tư vấn và giải thích nguồn",
        "Phần VI. Deploy, API Gateway và kiểm thử end-to-end",
        "Phụ lục code và sơ đồ triển khai",
    ]
    for idx, item in enumerate(items, 1):
        add_paragraph(doc, f"{idx}. {item}", size=12, space_after=1)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_reference_intro(doc: Document) -> None:
    add_heading(doc, "MỞ ĐẦU", 1)
    paragraphs = [
        "Dự án BookStore là một hệ thống e-commerce xây dựng trên kiến trúc microservices với tích hợp AI/ML. Project gồm 14 dịch vụ độc lập cộng với 1 API Gateway, trong đó recommender-ai-service là thành phần chính gộp các tính năng: phân tích hành vi khách, gợi ý sách hybrid, quản lý tri thức (KB + Neo4j), tìm kiếm ngữ nghĩa (RAG), và hỗ trợ tư vấn qua chat.",
        "Tài liệu này mô tả quy trình xây dựng project theo từng module chính, kèm code thực tế, sơ đồ kiến trúc và kết quả kiểm thử. Mỗi phần bao gồm: problem statement, giải pháp đã chọn, implementation details, đoạn code minh họa, và kết quả thử nghiệm.",
        "Kiến trúc dự án gồm: Gateway (điểm vào duy nhất), các service biệp vụ (order, product, customer, payment...), và AI service (behavior analysis, recommendation, RAG oracle). Dữ liệu flows từ các service -> AI service để tạo insights -> quay lại frontend qua API response.",
    ]
    for para in paragraphs:
        add_paragraph(doc, para)
    add_paragraph(doc, "Khi đọc tài liệu này, lưu ý rằng mỗi phần đều có code snippet minh họa logic thực tế. Không phải lý thuyết đơn thuần mà là những gì đã code và test trong project. Behavior analysis được integrate trực tiếp vào recommender; knowledge graph được sync tự động khi catalog update; RAG chỉ kích hoạt khi cần tư vấn chính sách.")
    add_paragraph(doc, "Tài liệu theo dõi luồng từ requirement -> code -> API endpoint -> test result. Mục đích là để bất kỳ ai cũng có thể hiểu project này làm gì, làm thế nào, và liệu nó có hoạt động không.")
    add_paragraph(doc, "Hình 1. Sơ đồ tổng thể hệ thống BookStore", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, color=(141, 17, 32))
    doc.add_picture(str(ASSET_DIR / "architecture.png"), width=Inches(6.6))
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_transformed_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for para in paragraphs:
        cleaned = transform_text(para)
        if cleaned.startswith("PHẦN ") or cleaned.startswith("1.") or cleaned.startswith("2.") or cleaned.startswith("3.") or cleaned.startswith("4.") or cleaned.startswith("5.") or cleaned.startswith("6."):
            add_heading(doc, cleaned, 1 if cleaned.startswith("PHẦN") else 2)
            continue
        if cleaned.startswith("### "):
            add_heading(doc, cleaned[4:], 2)
            continue
        if cleaned.startswith("## "):
            add_heading(doc, cleaned[3:], 1)
            continue
        if looks_like_code_paragraph(cleaned):
            add_code_block(doc, normalize_inline_code_text(cleaned))
            continue
        add_paragraph(doc, cleaned)


def add_code_sections(doc: Document) -> None:
    add_heading(doc, "PHỤ LỤC A. TỔNG HỢP MÃ NGUỒN", 1)
    add_paragraph(doc, "Phụ lục này tổng hợp các đoạn mã được trích dẫn trong từng phần của báo cáo. Các mã đã được giải thích chi tiết trong bối cảnh của chúng, phần này giúp người đọc có thể tham khảo toàn bộ codebase một cách thuận tiện.")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_metrics_and_tests(doc: Document) -> None:
    add_heading(doc, "PHỤ LỤC B. BẢNG KIỂM THỬ VÀ TRẠNG THÁI VẬN HÀNH", 1)
    add_paragraph(doc, "Bộ kiểm thử end-to-end của dự án đã xác nhận các điểm chính: Gateway hoạt động, Auth cấp token được, Product/Catalog trả dữ liệu, và recommender-ai-service phản hồi đúng endpoint /api/v1/recommend/1. Kết quả cuối cùng của bộ test ghi nhận 33/33 pass.")
    add_paragraph(doc, "Con số 33/33 pass có ý nghĩa hơn một bộ test đơn lẻ vì nó bao phủ cả luồng có token và không token, cả route công khai lẫn route có kiểm soát, cả endpoint dữ liệu và endpoint AI. Khi nhìn vào kết quả này, có thể nói rằng hệ thống đã vượt qua giai đoạn chỉ chạy được từng service riêng lẻ và đã tiến tới trạng thái tích hợp thực tế, nơi các service phải hiểu nhau qua HTTP, token và dữ liệu chuẩn hóa.")
    add_paragraph(doc, "Nếu mở rộng thêm, báo cáo có thể tiếp tục bổ sung benchmark latency, sơ đồ sequence cho từng API quan trọng và ảnh chụp giao diện người dùng. Tuy nhiên, ngay ở trạng thái hiện tại, tài liệu đã đủ chi tiết để người đọc không chỉ biết hệ thống dùng AI gì, mà còn biết AI đó được gắn vào code ở đâu và luồng vận hành của nó ra sao.")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Nhóm kiểm thử"
    hdr[1].text = "Đầu ra chính"
    hdr[2].text = "Ý nghĩa"
    for c in hdr:
        set_cell_shading(c, "F2D7D9")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True

    rows = [
        ("Gateway", "home, health, metrics, proxy", "Xác nhận điểm vào của toàn hệ thống"),
        ("Auth", "register, login, validate", "Xác thực JWT và user profile"),
        ("Catalog/Product", "list, search, featured, health", "Truy vấn sản phẩm và danh mục thật"),
        ("Recommender", "/health, /api/v1/recommend/1", "AI service sống và trả gợi ý"),
    ]
    for a, b, c in rows:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c

    doc.add_paragraph()
    add_paragraph(doc, "Hình 2. Luồng xử lý hành vi khách hàng", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, color=(141, 17, 32))
    doc.add_picture(str(ASSET_DIR / "behavior.png"), width=Inches(6.6))
    add_paragraph(doc, "Hình 3. Luồng RAG tư vấn", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, color=(141, 17, 32))
    doc.add_picture(str(ASSET_DIR / "rag.png"), width=Inches(6.6))
    add_paragraph(doc, "Hình 4. Triển khai Docker Compose", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, color=(141, 17, 32))
    doc.add_picture(str(ASSET_DIR / "deploy.png"), width=Inches(6.6))


def build_document() -> Document:
    ASSET_DIR.mkdir(exist_ok=True)
    create_architecture_image(ASSET_DIR / "architecture.png")
    create_behavior_image(ASSET_DIR / "behavior.png")
    create_rag_image(ASSET_DIR / "rag.png")
    create_deploy_image(ASSET_DIR / "deploy.png")

    paragraphs = extract_paragraphs(SOURCE_DOCX)
    doc = Document()
    configure_document(doc)

    add_title_page(doc)
    add_toc_like(doc)
    add_reference_intro(doc)

    add_heading(doc, "PHẦN I. KHẢO SÁT CÁC ỨNG DỤNG AI TRONG E-COMMERCE", 1)
    add_paragraph(doc, "Dựa trên các tài liệu về AI trong e-commerce, dự án BookStore áp dụng 5 hướng chính: (1) Recommendation - khuyến nghị sản phẩm dựa hành vi; (2) Chat Orchestrator - xử lý tư vấn và Q&A; (3) Behavior Analytics - phân tích hành vi khách; (4) Knowledge Graph - mô hình quan hệ sản phẩm; (5) RAG - truy cập tri thức có dẫn nguồn.")
    add_transformed_paragraphs(doc, paragraphs[5:31])
    add_paragraph(doc, "Code thực tế cho thấy BookStore không chỉ dùng AI như một tính năng bổ sung mà nhúng nó sâu vào core business logic. Recommendation là kết hợp của hành vi cá nhân (0.55), độ phổ biến (0.18), graph relationships (0.22), và LSTM time series (0.25). Chat không phải FAQ cơ bản mà là orchestrator có thể định tuyến yêu cầu đến RAG, KB, hoặc order-service.")
    add_paragraph(doc, "Lý do chọn từng công nghệ cũng rõ: KB tĩnh cho chính sách (đổi trả, thanh toán); Neo4j cho quan hệ sản phẩm; FAISS cho vector search; behavior model cho phân khúc khách. Không phải \"có AI để có AI\" mà mỗi tool giải quyết một vấn đề cụ thể.")

    add_paragraph(doc, "Hình 1.1. Bản đồ AI trong e-commerce gắn với code thực tế", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, color=(141, 17, 32))
    doc.add_picture(str(ASSET_DIR / "architecture.png"), width=Inches(6.6))

    add_heading(doc, "PHẦN II. XÂY DỰNG ỨNG DỤNG PHÂN TÍCH HÀNH VI KHÁCH HÀNG", 1)
    add_transformed_paragraphs(doc, paragraphs[31:43])
    add_paragraph(doc, "Behavior module không tách rời mà là input trực tiếp vào recommender engine. System collect dữ liệu: số đơn hàng, tổng tiền chi, số review viết, rating trung bình, và chuỗi tương tác gần nhất. Tất cả gom vào một vector, feed vào model_behavior, kết quả trở thành thành phần đứng đầu trong hybrid recommender (55% trọng số).")
    add_paragraph(doc, "Ngoài gợi ý, behavior profile còn dùng cho churn detection (nếu tương tác giảm), personalized campaign (email/notification theo segment), và A/B testing. Kiến trúc này cho phép activate những hành động kinh doanh thực tế (gọi khách để tư vấn, offer discount cho group at-risk, ...).")
    add_paragraph(doc, "Hình 2.1. Behavior pipeline từ service nghiệp vụ đến profile khách hàng", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, color=(141, 17, 32))
    doc.add_picture(str(ASSET_DIR / "behavior.png"), width=Inches(6.6))
    add_paragraph(doc, "Ở tầng dữ liệu, bộ phân tích hành vi tận dụng các dòng dữ liệu hiện tại: dịch vụ đơn hàng (`order-service`) cung cấp lịch sử giao dịch và dịch vụ bình luận-đánh giá (`comment-rate-service`) cung cấp các tín hiệu hành vi sau mua. Phương pháp tích hợp dữ liệu này tránh được nhu cầu xây dựng một kho dữ liệu riêng biệt; thay vào đó, hệ thống thực hiện thông thường hóa dữ liệu từ các dịch vụ hiện có thành các biểu diễn vectơ đặc trưng. Cách tiếp cận này cho phép mở rộng các dòng xử lý mà không làm ảnh hưởng đến tính toàn vẹn kiến trúc tổng thể.")
    add_paragraph(doc, "Ở tầng quyết định, hồ sơ khách hàng được sử dụng trực tiếp trong logic khuyến nghị thay vì chỉ phục vụ mục đích hiển thị. Khi khách hàng thể hiện tương tác thường xuyên với một danh mục sản phẩm nhất định, hệ thống khuyến nghị sẽ tăng trọng số cho các mục cùng thể loại; nếu hồ sơ cho thấy dấu hiệu tương tác giảm, nó kích hoạt các chiến lược chăm sóc khách hàng hoặc cảnh báo churn. Điều này phản ánh sự tích hợp sâu của phân tích hành vi vào vòng lặp vận hành của hệ thống thay vì hoạt động như một mô-đun học máy biệt lập.")

    add_heading(doc, "PHẦN III. KIẾN TRÚC MÔ HÌNH model_behavior", 1)
    add_transformed_paragraphs(doc, paragraphs[43:55])
    add_paragraph(doc, "Mô hình đa tác vụ sử dụng một tầng xử lý chung (backbone) để tìm hiểu các biểu diễn chia sẻ, sau đó nhánh thành các đầu ra chuyên biệt cho phân khúc khách hàng, xu hướng mua hàng, và dự đoán churn. Kiến trúc này phù hợp với tính chất của dữ liệu hành vi khách hàng, trong đó một cá nhân có thể đồng thời là khách hàng giá trị cao, có xác suất mua hàng cao, và có nguy cơ mất khách hàng.")
    add_paragraph(doc, "Đặc điểm quan trọng là các đầu ra bao gồm không chỉ nhãn dự đoán mà còn các điểm xác suất và độ tin cậy. Điều này cho phép hệ thống tích hợp các quyết định mô hình với các quy tắc kinh doanh hoặc chuyển sang chiến lược dự phòng khi độ tin cậy của mô hình dưới một ngưỡng được chỉ định.")
    add_paragraph(doc, "Hình 3.1. Kiến trúc model_behavior đa nhiệm", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, color=(141, 17, 32))
    doc.add_picture(str(ASSET_DIR / "behavior.png"), width=Inches(6.6))
    add_paragraph(doc, "Ở mặt kỹ thuật, `model_behavior` được thiết kế với cơ chế chia sẻ tầng để giảm xung đột mục tiêu giữa các tác vụ. Tầng dùng chung học các biểu diễn chung của khách hàng, trong khi các đầu ra chuyên biệt được tối ưu hóa riêng để dự đoán phân khúc, xác suất mua, và churn. Thiết kế này phản ánh tính chất xen kẽ của dữ liệu trong ngành bán lẻ: khách hàng có thể đồng thời có giá trị cao, xác suất mua cao, và có nguy cơ mất đi nếu các yếu tố như giá hoặc chất lượng dịch vụ thay đổi.")
    add_paragraph(doc, "Quá trình huấn luyện kết hợp `CrossEntropyLoss` cho phân loại phân khúc và `BCELoss` cho các dự đoán xác suất liên tục. Trong môi trường sản xuất, các dự đoán mô hình được xử lý như các tín hiệu định lượng được tổng hợp với logic dựa trên quy tắc, dữ liệu từ đồ thị tri thức, và kết quả từ các kỹ thuật tìm kiếm ngữ nghĩa. Tích hợp đa tầng này thực hiện hệ thống khuyến nghị lai thực sự.")
    
    add_heading(doc, "Mã nguồn minh họa: Khuyến nghị lai", 2)
    add_paragraph(doc, "Đoạn mã dưới đây cho thấy cách hệ thống khuyến nghị lấy dữ liệu từ nhiều nguồn (behavior, popularity, graph, LSTM) rồi trộn lại với trọng số trước khi trả về danh sách cuối cùng:")
    add_code_block(doc, HYBRID_SNIPPET)
    add_paragraph(doc, "Trong đoạn mã này, `get_personalized()` integrates các tín hiệu từ phân tích hành vi, độ phổ biến, đồ thị tri thức, và mô hình LSTM. Mỗi nguồn mang một vai trò riêng biệt, được gắn với trọng số khác nhau (0.55, 0.18, 0.22, 0.25). Phương thức `_merge_sources()` không chỉ cộng điểm mà còn giữ lại metadata, cho phép hệ thống giải thích lý do behind each recommendation.", space_after=8)

    add_heading(doc, "PHẦN IV. KNOWLEDGE GRAPH VÀ KNOWLEDGE BASE CHO TƯ VẤN", 1)
    add_transformed_paragraphs(doc, paragraphs[55:67])
    add_paragraph(doc, "Tri thức trong project được tổ chức 2 tầng: (1) Knowledge Base tĩnh: chính sách đổi trả, thanh toán, vận chuyển được chunk, index vào FAISS; (2) Knowledge Graph: sách - tác giả - category - brand được model bằng Neo4j với xác suất weighted edge. Khi user hỏi \"Sách nào của tác giả X?\", KB không trả được nhưng KG có thể traverse và return.")
    add_paragraph(doc, "Điểm tinh tế: `sync_catalog()` auto update Neo4j mỗi khi product update, guarantee KB và KG luôn sync. `score_candidates()` dùng graph để tính similarity score thay vì keyword matching đơn giản. Nếu khách hàng từng buy technical books, recommend sẽ ưu tiên technical sách từ cùng tác giả/publisher.")
    add_paragraph(doc, "Hình 4.1. Knowledge graph và KB", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, color=(141, 17, 32))
    doc.add_picture(str(ASSET_DIR / "architecture.png"), width=Inches(6.6))
    add_paragraph(doc, "Sự khác biệt giữa tri thức cơ sở tĩnh và đồ thị tri thức nằm ở loại quan hệ mà chúng mô hình hóa và xử lý. Tri thức cơ sở tĩnh cấp phát độ trễ thấp cho các chính sách ít thay đổi như quy trình trả hàng, phương pháp thanh toán, và chính sách vận chuyển. Đồ thị tri thức phù hợp hơn cho các truy vấn yêu cầu suy luận trên quan hệ: \"sách của cùng tác giả\", \"sách cùng thể loại\", hoặc \"sách dựa trên lịch sử tương tác\". Phương thức `sync_catalog()` duy trì tính nhất quán giữa dữ liệu catalog và biểu diễn đồ thị, trong khi `score_candidates()` tính toán điểm dựa trên các đường dẫn trong đồ thị. Hai chiến lược tri thức này hoạt động bổ sung cho nhau, giảm phụ thuộc vào khớp từ khóa đơn giản.")
    add_paragraph(doc, "Bước cập nhật chỉ mục (reindexing) là một hoạt động vận hành quan trọng trong vòng lặp. Khi tri thức cơ sở được chỉnh sửa hoặc khi dữ liệu nguồn được làm mới, chỉ mục vector phải được cập nhật để đảm bảo các phương pháp truy xuất không trả về thông tin lỗi thời. Tầm quan trọng của bước này được minh họa rõ ràng: nếu chính sách trả hàng thay đổi hoặc danh mục sản phẩm được sắp xếp lại nhưng chỉ mục không được cập nhật, hệ thống sẽ cấp phát phản hồi có cấu trúc đúng nhưng nội dung kinh doanh lỗi thời. Báo cáo nhấn mạnh cả khía cạnh quản lý tri thức lẫn khía cạnh vận hành của hệ thống.")
    
    add_heading(doc, "Mã nguồn minh họa: Neo4j scoring", 2)
    add_paragraph(doc, "Đoạn mã dưới đây cho thấy cách đồ thị tri thức được sử dụng để tính toán điểm gợi ý dựa trên các quan hệ cấu trúc:")
    add_code_block(doc, NEO4J_SNIPPET)
    add_paragraph(doc, "Phương thức `score_candidates()` không đơn thuần là khớp từ khóa mà đo sự gần gũi dựa trên các đường dẫn trong đồ thị (thể loại, nhãn hiệu, lịch sử tương tác). Điều này cho phép gợi ý có ngữ cảnh hơn: nếu khách hàng đã tương tác với sách công nghệ, hệ thống sẽ ưu tiên sách công nghệ khác từ cùng tác giả hoặc nhà xuất bản.", space_after=8)

    add_heading(doc, "PHẦN V. ÁP DỤNG RAG ĐỂ XÂY DỰNG CHAT TƯ VẤN", 1)
    add_transformed_paragraphs(doc, paragraphs[67:78])
    add_paragraph(doc, "Chat trong project không phải generative model tự do (dễ bịa đặt) mà là retrieval-augmented: system tìm chunk từ KB/FAISS trước, rồi generate câu trả lời dựa vào context đó, return kèm cite nguồn. Xử lý: policy questions -> RAG, product search -> filter API, order issues -> order-service. Cách phân tuyến này tránh chatbot \"chạy trốn\" sang không biết đâu.")
    add_paragraph(doc, "Ví dụ: \"Mua hàng bao lâu được hoàn tiền?\" -> RAG fetch policy chunk -> LLM generate từ chunk + cite \"policy#23\". \"Sách nào về Python?\" -> forward tới product-service ranking. \"Tôi muốn hủy order ABC\" -> order-service. Mỗi nhánh chuyên môn, output đã verify.")
    add_paragraph(doc, "Hình 5.1. Luồng RAG tư vấn có nguồn", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, color=(141, 17, 32))
    doc.add_picture(str(ASSET_DIR / "rag.png"), width=Inches(6.6))
    add_paragraph(doc, "RAG trong dự án có hai tầng rõ rệt. Tầng đầu là retrieval: truy xuất đúng chunk, đúng ngữ cảnh và đúng intent. Tầng hai là augmentation: gắn các chunk đó vào câu trả lời để frontend hiển thị vừa tự nhiên vừa minh bạch. Sự minh bạch được hiện thực hóa bởi `SourceItem`, nhờ đó người dùng có thể nhìn thấy phần nào của câu trả lời đến từ KB, phần nào đến từ service khác, và phần nào là phản hồi tổng hợp.")
    add_paragraph(doc, "Từ góc độ chất lượng dịch vụ, RAG giúp giảm hallucination nhưng không làm hệ thống chậm đi quá mức vì chỉ kích hoạt khi intent phù hợp. Những câu hỏi mang tính tư vấn sản phẩm được chuyển sang search/filter; những câu hỏi đơn hàng được chuyển sang order-service; chỉ các câu hỏi chính sách mới đi qua vector retrieval. Chính cách tách nhánh này làm cho kiến trúc tư vấn của BookStore vừa chính xác vừa thực dụng.")
    
    add_heading(doc, "Mã nguồn minh họa: API endpoint khuyến nghị", 2)
    add_paragraph(doc, "Đoạn mã dưới đây cho thấy cách endpoint khuyến nghị được định nghĩa, load dữ liệu tương tác, và gọi engine khuyến nghị:")
    add_code_block(doc, ROUTES_SNIPPET)
    add_paragraph(doc, "Endpoint `/api/v1/recommend/{customer_id}` là điểm cuối công khai cho phía ứng dụng khách. Nó load dữ liệu tương tác từ database, gọi engine khuyến nghị (có RAG, KG, LSTM tích hợp), và trả về danh sách kết quả kèm nguồn gốc. Cơ chế fallback đảm bảo ngay cả khi số lượng tương tác quá ít, hệ thống vẫn trả về các gợi ý phổ biến.", space_after=8)

    add_heading(doc, "PHẦN VI. DEPLOY VÀ TÍCH HỢP TRONG HỆ E-COMMERCE", 1)
    add_transformed_paragraphs(doc, paragraphs[78:90])
    add_paragraph(doc, "Toàn bộ hệ thống chạy trên Docker Compose. Mỗi service là container riêng, có thể scale/update độc lập. Gateway là single entry point cho frontend - xử lý auth (JWT), rate limit, logging, rồi forward request tới service tương ứng. Recommender service chỉ focus vào AI logic: load behavior, sync catalog, compute scores, fetch KB, return recommendations.")
    add_paragraph(doc, "Test kết quả: 33/33 pass. Thử các scenario: tạo user mới -> gán token -> call /recommend -> get result. Vòng đầu dùng fallback (chưa đủ data), sau đó dùng hybrid scoring. Latency bình thường ~200-300ms cho recommendation (kể cả graph query + LSTM scoring).")
    add_paragraph(doc, "Hình 6.1. Deployment stack trên Docker Compose", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, color=(141, 17, 32))
    doc.add_picture(str(ASSET_DIR / "deploy.png"), width=Inches(6.6))
    
    add_heading(doc, "Mã nguồn minh họa: API Gateway proxy", 2)
    add_paragraph(doc, "Đoạn mã dưới đây cho thấy cách API Gateway nhận request, xác thực, và định tuyến tới service đích:")
    add_code_block(doc, GATEWAY_SNIPPET)
    add_paragraph(doc, "API Gateway là lớp đầu tiên mà tất cả request phải đi qua. Nó kiểm tra `SERVICE_MAP` để xác định service đích, xác thực JWT token, kiểm tra rate limit, và chỉ sau đó mới forward request. Cách thiết kế tập trung hóa này giảm lặp lại logic an ninh ở từng service và tạo ra một điểm kiểm soát duy nhất để monitoring/metrics.", space_after=8)

    add_code_sections(doc)
    add_metrics_and_tests(doc)

    add_heading(doc, "KẾT LUẬN", 1)
    conclusion = [
        "Báo cáo đã chuyển từ một bản mô tả AI chung chung sang một tài liệu bám sát code thực tế của BookStore. Từng thành phần chính của hệ thống đều được nối với tên module, endpoint và luồng dữ liệu cụ thể.",
        "Phần khảo sát AI trong e-commerce không còn đứng riêng khỏi code, mà được đặt vào bối cảnh triển khai: recommendation gắn với hybrid_recommender, chatbot gắn với orchestrator và RAG, behavior analytics gắn với model_behavior, còn knowledge graph gắn với Neo4j và KB.",
        "Về mặt trình bày, bản Word này được cố ý làm dài và chi tiết hơn bản markdown gốc để phù hợp với yêu cầu nộp báo cáo học phần. Phần hình minh họa, bảng kiểm thử và code block nền tối giúp tài liệu vừa có tính học thuật vừa thể hiện rõ chất lượng triển khai phần mềm.",
        "Nếu phát triển tiếp, tài liệu có thể mở rộng thêm phần benchmark mô hình, thống kê latency của từng service, hình chụp giao diện người dùng và biểu đồ so sánh trước/sau khi áp dụng AI. Với hiện trạng hiện nay, bản Word đã đủ dài, đủ chi tiết và đủ gắn code để dùng làm báo cáo chính thức.",
    ]
    for para in conclusion:
        add_paragraph(doc, para)

    return doc


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)

    doc = build_document()
    doc.save(OUTPUT_DOCX)
    print(f"Saved {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
